from docx import Document
import sys

def extract_docx(filepath, output_path):
    doc = Document(filepath)
    content = []
    
    for para in doc.paragraphs:
        content.append(para.text)
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text for cell in row.cells])
            content.append(row_text)
    
    text = "\n".join(content)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    
    print(f"Extracted {len(content)} paragraphs to {output_path}")
    print("\n--- First 5000 characters ---\n")
    print(text[:5000])

if __name__ == "__main__":
    extract_docx(
        r"C:\Users\Admin\.gemini\antigravity\scratch\BÁO CÁO THỰC TẬP TỐT NGHIỆP - LƯ LỆ VÂN.docx",
        r"C:\Users\Admin\.gemini\antigravity\scratch\sample_thesis_content.txt"
    )
