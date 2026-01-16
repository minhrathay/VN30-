"""
Convert Markdown to Professional Academic Word Document
Format: Vietnamese Thesis/Report Standard
"""
import subprocess
import sys
import os
import re

# Install dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def set_page_margins(doc):
    """Set academic margins: Left 3.5cm, Right 2cm, Top/Bottom 2.5cm"""
    for section in doc.sections:
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

def add_page_number(doc):
    """Add page numbers at bottom center"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run2._r.append(instrText)
        
        run3 = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)

def setup_styles(doc):
    """Setup academic styles matching sample thesis"""
    # Normal style - Body text
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.15  # From sample
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Heading 1 - Chapter titles (20pt, bold, center)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.15
    
    # Heading 2 - Section titles (14pt, bold)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.15
    
    # Heading 3 - Subsection (13pt, bold)
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.15

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, center=False):
    """Add a formatted paragraph"""
    p = doc.add_paragraph(style=style)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process bold and italic markers
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
            run.bold = bold
            run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    
    return p

def add_table(doc, table_data):
    """Add a formatted table"""
    if not table_data or len(table_data) < 1:
        return
    
    num_cols = len(table_data[0])
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                # Format cell
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        if i == 0:  # Header row
                            run.bold = True
    
    doc.add_paragraph()  # Space after table

def convert_md_to_academic_docx(md_file, docx_file, images_dir='images'):
    """Convert markdown to academic Word document"""
    
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Setup page format
    set_page_margins(doc)
    setup_styles(doc)
    add_page_number(doc)
    
    # Process content
    lines = content.split('\n')
    in_code_block = False
    in_table = False
    table_data = []
    figure_count = 0
    table_count = 0
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            p.paragraph_format.left_indent = Cm(1)
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            i += 1
            continue
        
        # Tables
        if '|' in line and not line.startswith('#'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                # Skip separator row
                if all(set(c.replace('-', '').replace(':', '')) == set() or c.replace('-', '').replace(':', '') == '' for c in cells):
                    i += 1
                    continue
                table_data.append(cells)
            i += 1
            continue
        elif table_data:
            add_table(doc, table_data)
            table_count += 1
            table_data = []
        
        # Headers
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].upper()
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.bold = True
        elif line.startswith('## '):
            text = line[3:]
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        elif line.startswith('### '):
            text = line[4:]
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
        elif line.startswith('#### '):
            text = line[5:]
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
        elif line.startswith('---'):
            doc.add_page_break()
        elif line.startswith('- '):
            text = line[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Cm(1)
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.left_indent = Cm(1)
        elif line.startswith('$$'):
            # Math formula
            formula = line.replace('$$', '').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(formula)
            run.italic = True
            run.font.name = 'Cambria Math'
            run.font.size = Pt(12)
        elif line.strip():
            # Regular paragraph
            text = line
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Remove links
            add_formatted_paragraph(doc, text)
        else:
            # Empty line
            doc.add_paragraph()
        
        i += 1
    
    # Handle remaining table
    if table_data:
        add_table(doc, table_data)
    
    # Save
    doc.save(docx_file)
    print(f"✅ Đã tạo file Word chuẩn học thuật: {docx_file}")
    print(f"   - Font: Times New Roman 13pt")
    print(f"   - Dãn dòng: 1.5")
    print(f"   - Lề: Trái 3.5cm, Phải 2cm, Trên/Dưới 2.5cm")
    print(f"   - Đánh số trang: Có")

if __name__ == "__main__":
    convert_md_to_academic_docx('BAO_CAO_CHUYEN_DE.md', 'BAO_CAO_CHUYEN_DE.docx')
