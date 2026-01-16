"""Extract formatting details from sample Word document"""
from docx import Document
from docx.shared import Pt, Cm, Twips

def extract_formatting(filepath):
    doc = Document(filepath)
    
    print("=" * 60)
    print("FORMATTING ANALYSIS - Sample Thesis")
    print("=" * 60)
    
    # Page margins
    print("\n📐 PAGE MARGINS:")
    for i, section in enumerate(doc.sections):
        print(f"  Left:   {section.left_margin.cm:.2f} cm")
        print(f"  Right:  {section.right_margin.cm:.2f} cm")
        print(f"  Top:    {section.top_margin.cm:.2f} cm")
        print(f"  Bottom: {section.bottom_margin.cm:.2f} cm")
        print(f"  Page width:  {section.page_width.cm:.2f} cm")
        print(f"  Page height: {section.page_height.cm:.2f} cm")
        break  # Only first section
    
    # Font styles from paragraphs
    print("\n📝 FONT ANALYSIS (from first 30 paragraphs):")
    fonts_found = {}
    sizes_found = {}
    
    for para in doc.paragraphs[:30]:
        for run in para.runs:
            if run.font.name:
                fonts_found[run.font.name] = fonts_found.get(run.font.name, 0) + 1
            if run.font.size:
                size_pt = run.font.size.pt
                sizes_found[size_pt] = sizes_found.get(size_pt, 0) + 1
    
    print("  Fonts used:")
    for font, count in sorted(fonts_found.items(), key=lambda x: -x[1]):
        print(f"    - {font}: {count} occurrences")
    
    print("  Font sizes used:")
    for size, count in sorted(sizes_found.items(), key=lambda x: -x[1]):
        print(f"    - {size:.1f} pt: {count} occurrences")
    
    # Paragraph formatting
    print("\n📏 PARAGRAPH FORMATTING:")
    for para in doc.paragraphs[:10]:
        if para.text.strip():
            pf = para.paragraph_format
            print(f"  Paragraph: '{para.text[:40]}...'")
            if pf.line_spacing:
                print(f"    Line spacing: {pf.line_spacing}")
            if pf.first_line_indent:
                print(f"    First line indent: {pf.first_line_indent.cm:.2f} cm")
            if pf.left_indent:
                print(f"    Left indent: {pf.left_indent.cm:.2f} cm")
            break
    
    # Styles
    print("\n🎨 STYLES USED:")
    styles_used = set()
    for para in doc.paragraphs:
        if para.style:
            styles_used.add(para.style.name)
    for style in sorted(styles_used):
        print(f"  - {style}")

if __name__ == "__main__":
    extract_formatting(r"C:\Users\Admin\.gemini\antigravity\scratch\BÁO CÁO THỰC TẬP TỐT NGHIỆP - LƯ LỆ VÂN.docx")
